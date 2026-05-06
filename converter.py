# converter.py
import os
from markdown import markdown
from weasyprint import HTML
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.oxml.ns import qn
from themes import THEMES


def convert_md_to_pdf(md_path, output_dir, theme_name):
    """Convert a Markdown file to PDF with selected theme."""
    with open(md_path, 'r', encoding='utf-8') as f:
        md_text = f.read()

    html_content = markdown(md_text, extensions=['tables', 'fenced_code', 'codehilite', 'toc'])

    css_string = THEMES.get(theme_name, THEMES['Default'])

    full_html = f"""<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<style>
{css_string}
</style>
</head>
<body>
{html_content}
</body>
</html>"""

    os.makedirs(output_dir, exist_ok=True)

    base_name = os.path.splitext(os.path.basename(md_path))[0]
    pdf_path = os.path.join(output_dir, base_name + '.pdf')

    HTML(string=full_html).write_pdf(pdf_path)

    return pdf_path


def convert_md_to_docx(md_path, output_dir):
    """Convert a Markdown file to DOCX with Default-theme-like styling."""
    with open(md_path, 'r', encoding='utf-8') as f:
        md_text = f.read()

    html_content = markdown(md_text, extensions=['tables', 'fenced_code', 'codehilite', 'toc'])

    os.makedirs(output_dir, exist_ok=True)
    base_name = os.path.splitext(os.path.basename(md_path))[0]
    docx_path = os.path.join(output_dir, base_name + '.docx')

    doc = Document()

    default_style = doc.styles['Normal']
    font = default_style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    pf = default_style.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.15

    soup = BeautifulSoup(html_content, 'html.parser')

    def _add_inline_content(tag, paragraph):
        """Recursively add inline content from a tag to a paragraph."""
        for child in tag.children:
            if isinstance(child, str):
                paragraph.add_run(child)
            elif child.name in ('strong', 'b'):
                run = paragraph.add_run()
                run.bold = True
                _inline_to_run(child, run)
            elif child.name in ('em', 'i'):
                run = paragraph.add_run()
                run.italic = True
                _inline_to_run(child, run)
            elif child.name == 'code':
                run = paragraph.add_run(child.get_text())
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            elif child.name == 'a':
                run = paragraph.add_run(child.get_text())
                run.font.color.rgb = RGBColor(0x2E, 0x86, 0xAB)
                run.font.underline = True
            elif child.name == 'br':
                paragraph.add_run('\n')
            elif child.name:
                run = paragraph.add_run()
                _inline_to_run(child, run)

    def _inline_to_run(tag, run):
        """Recursively add inline element content to a run."""
        for child in tag.children:
            if isinstance(child, str):
                run.text += child
            elif child.name in ('strong', 'b'):
                run.bold = True
                _inline_to_run(child, run)
            elif child.name in ('em', 'i'):
                run.italic = True
                _inline_to_run(child, run)
            elif child.name == 'code':
                run.text += child.get_text()
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
            else:
                _inline_to_run(child, run)

    def _add_table(table_tag, doc):
        rows = table_tag.find_all('tr')
        if not rows:
            return

        num_rows = len(rows)
        num_cols = 0
        for row in rows:
            cells = row.find_all(['th', 'td'])
            num_cols = max(num_cols, len(cells))

        if num_cols == 0:
            return

        table = doc.add_table(rows=num_rows, cols=num_cols)
        table.style = 'Table Grid'

        for i, row in enumerate(rows):
            cells = row.find_all(['th', 'td'])
            for j, cell in enumerate(cells):
                if j >= num_cols:
                    break
                cell_obj = table.cell(i, j)
                cell_obj.text = ''
                para = cell_obj.paragraphs[0]
                _add_inline_content(cell, para)
                if i == 0:
                    for run in para.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    shading_elm = cell_obj._element.get_or_add_tcPr()
                    shading = shading_elm.makeelement(qn('w:shd'), {
                        qn('w:fill'): '2E86AB',
                        qn('w:val'): 'clear'
                    })
                    shading_elm.append(shading)

    def _add_element(tag, doc):
        tag_name = tag.name

        if tag_name in ('h1', 'h2', 'h3', 'h4', 'h5', 'h6'):
            level = int(tag_name[1])
            heading = doc.add_heading('', level=level if level <= 3 else 3)
            heading.clear()
            if level == 1:
                for run_el in heading.runs:
                    run_el.font.size = Pt(20)
                    run_el.font.color.rgb = RGBColor(0x2E, 0x86, 0xAB)
                    run_el.bold = True
            elif level == 2:
                for run_el in heading.runs:
                    run_el.font.size = Pt(16)
                    run_el.font.color.rgb = RGBColor(0x2E, 0x86, 0xAB)
                    run_el.bold = True
            elif level == 3:
                for run_el in heading.runs:
                    run_el.font.size = Pt(13)
                    run_el.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
                    run_el.bold = True
            _add_inline_content(tag, heading)
            return

        if tag_name == 'pre':
            code_block = tag.find('code')
            code_text = code_block.get_text() if code_block else tag.get_text()
            para = doc.add_paragraph()
            pf = para.paragraph_format
            pf.space_before = Pt(4)
            pf.space_after = Pt(4)
            shading_elm = para._element.get_or_add_pPr()
            shading = shading_elm.makeelement(qn('w:shd'), {
                qn('w:fill'): 'F2F2F2',
                qn('w:val'): 'clear'
            })
            shading_elm.append(shading)
            lines = code_text.splitlines()
            for idx, line in enumerate(lines):
                run = para.add_run(line)
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
                if idx < len(lines) - 1:
                    run = para.add_run('\n')
                    run.font.name = 'Courier New'
                    run.font.size = Pt(10)
            return

        if tag_name == 'code' and tag.parent and tag.parent.name != 'pre':
            para = doc.add_paragraph()
            run = para.add_run(tag.get_text())
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            shading_elm = para._element.get_or_add_pPr()
            shading = shading_elm.makeelement(qn('w:shd'), {
                qn('w:fill'): 'F2F2F2',
                qn('w:val'): 'clear'
            })
            shading_elm.append(shading)
            return

        if tag_name == 'blockquote':
            for child in tag.children:
                if hasattr(child, 'name') and child.name:
                    _add_element(child, doc)
                elif str(child).strip():
                    para = doc.add_paragraph()
                    para.paragraph_format.left_indent = Cm(1)
                    run = para.add_run(str(child).strip())
                    run.italic = True
                    run.font.size = Pt(11)
            return

        if tag_name == 'table':
            _add_table(tag, doc)
            return

        if tag_name in ('p', 'div'):
            if not tag.get_text().strip() and not tag.find():
                return
            para = doc.add_paragraph()
            _add_inline_content(tag, para)
            return

        if tag_name in ('ul', 'ol'):
            for child in tag.children:
                if hasattr(child, 'name') and child.name == 'li':
                    para = doc.add_paragraph()
                    para.style = doc.styles['List Bullet']
                    pf = para.paragraph_format
                    pf.left_indent = Cm(1.27)
                    _add_inline_content(child, para)
            return

        if tag_name == 'hr':
            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = Pt(6)
            pPr = para._element.get_or_add_pPr()
            pBdr = pPr.makeelement(qn('w:pBdr'), {})
            bottom = pBdr.makeelement(qn('w:bottom'), {
                qn('w:val'): 'single',
                qn('w:sz'): '12',
                qn('w:space'): '1',
                qn('w:color'): 'CCCCCC'
            })
            pBdr.append(bottom)
            pPr.append(pBdr)
            return

        if tag_name == 'img':
            src = tag.get('src', '')
            alt = tag.get('alt', '')
            para = doc.add_paragraph()
            if src:
                try:
                    if not os.path.isabs(src):
                        md_dir = os.path.dirname(md_path)
                        src = os.path.join(md_dir, src)
                    if os.path.exists(src):
                        run = para.add_run('')
                        run.add_picture(src, width=Inches(4))
                    else:
                        run = para.add_run(f'[Image not found: {alt}]')
                        run.italic = True
                except Exception:
                    run = para.add_run(f'[Image: {alt}]')
                    run.italic = True
            return

    for child in soup.children:
        if hasattr(child, 'name') and child.name:
            _add_element(child, doc)

    doc.save(docx_path)
    return docx_path
